#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Markdown(경량 방언) -> DOCX 변환 빌더.

지원 문법
  # / ## / ### / ####      : 제목 1~4
  |a|b|  +  |---|---|      : 표 (구분선 위 행이 헤더)
  - , * , 두 칸 들여쓰기    : 글머리표(2단계)
  1.                       : 번호 목록
  > text                   : 강조 박스(음영 1셀 표)
  ---                      : 페이지 나눔
  [[TOC]]                  : 목차 필드
  [[TITLE]] ... 블록       : 표지
  **bold** , `code`        : 인라인 서식
"""
import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

BODY_FONT = "Noto Sans CJK KR"
MONO_FONT = "NanumGothicCoding"
ACCENT = RGBColor(0x0B, 0x4F, 0x6C)
TEXT_WIDTH = Cm(16.4)   # A4 210mm - 좌우 여백 22mm
MUTED = RGBColor(0x55, 0x5D, 0x66)


def set_run_font(run, name=BODY_FONT, size=None, bold=None, color=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`]+`)")


def add_inline(par, text, size=10.5, base_bold=False, color=None):
    for chunk in INLINE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            r = par.add_run(chunk[2:-2])
            set_run_font(r, BODY_FONT, size, True, color)
        elif chunk.startswith("*") and chunk.endswith("*") and len(chunk) > 2:
            r = par.add_run(chunk[1:-1])
            set_run_font(r, BODY_FONT, size, base_bold, color, italic=True)
        elif chunk.startswith("`") and chunk.endswith("`"):
            r = par.add_run(chunk[1:-1])
            set_run_font(r, MONO_FONT, size - 0.5, base_bold, RGBColor(0xA3, 0x1D, 0x1D))
        else:
            r = par.add_run(chunk)
            set_run_font(r, BODY_FONT, size, base_bold, color)


def shade(cell, hexcolor):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcpr.append(shd)


def cell_borders(cell, color="C6CDD4", sz=4):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:color"), color)
        borders.append(e)
    tcpr.append(borders)


def add_toc(doc):
    par = doc.add_paragraph()
    run = par.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "목차를 갱신하려면 이 영역을 선택 후 F9 를 누르십시오."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (fld, instr, sep, placeholder, end):
        run._r.append(el)


def add_page_number_footer(section):
    footer = section.footer
    par = footer.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run()
    set_run_font(run, BODY_FONT, 9, color=MUTED)
    for kind, txt in (("begin", None), (None, "PAGE"), ("end", None)):
        if kind:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), kind)
            run._r.append(fc)
        else:
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = txt
            run._r.append(it)


def style_doc(doc):
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.5)
    set_run_font_style(normal, BODY_FONT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.32
    specs = {
        "Heading 1": (17, True, ACCENT, 20, 10),
        "Heading 2": (13.5, True, ACCENT, 14, 6),
        "Heading 3": (11.5, True, RGBColor(0x1F, 0x2A, 0x33), 10, 4),
        "Heading 4": (10.5, True, RGBColor(0x33, 0x40, 0x4A), 8, 3),
    }
    for name, (size, bold, color, before, after) in specs.items():
        st = doc.styles[name]
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = color
        set_run_font_style(st, BODY_FONT)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def set_run_font_style(style, name):
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)


def build(src_path, out_path):
    lines = open(src_path, encoding="utf-8").read().split("\n")
    doc = Document()
    style_doc(doc)
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.2)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    add_page_number_footer(sec)

    i = 0
    n = len(lines)
    emitted_break = True  # title page starts fresh
    while i < n:
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped == "":
            i += 1
            continue

        if stripped not in ("[[PAGEBREAK]]",) and not stripped.startswith("#"):
            emitted_break = False

        if stripped == "[[PAGEBREAK]]":
            if not emitted_break:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                emitted_break = True
            i += 1
            continue

        if stripped == "[[TOC]]":
            add_toc(doc)
            i += 1
            continue

        if stripped.startswith("[[TITLE]]"):
            title = stripped[len("[[TITLE]]"):].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(150)
            set_run_font(p.add_run(title), BODY_FONT, 26, True, ACCENT)
            i += 1
            continue

        if stripped.startswith("[[SECTIONTITLE]]"):
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(12)
            set_run_font(p2.add_run(stripped[len("[[SECTIONTITLE]]"):].strip()),
                         BODY_FONT, 17, True, ACCENT)
            i += 1
            continue

        if stripped.startswith("[[SUB]]"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.add_run(stripped[len("[[SUB]]"):].strip()), BODY_FONT, 13, False, MUTED)
            i += 1
            continue

        if stripped.startswith("[[META]]"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            set_run_font(p.add_run(stripped[len("[[META]]"):].strip()), BODY_FONT, 9.5, False, MUTED)
            i += 1
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            if level == 1 and not emitted_break:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            h = doc.add_heading(level=level)
            add_inline(h, m.group(2), {1: 17, 2: 13.5, 3: 11.5, 4: 10.5}[level], True,
                       ACCENT if level <= 2 else RGBColor(0x1F, 0x2A, 0x33))
            emitted_break = False
            i += 1
            continue

        m = re.match(r"^!\[(.*)\]\((.+?)\)$", stripped)
        if m:
            cap, path = m.group(1), m.group(2)
            if not os.path.exists(path):
                raise SystemExit(f"figure not found: {path}")
            par = doc.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            par.paragraph_format.space_before = Pt(8)
            par.paragraph_format.space_after = Pt(2)
            par.paragraph_format.keep_with_next = True
            run = par.add_run()
            run.add_picture(path, width=TEXT_WIDTH)
            if cap:
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_after = Pt(11)
                num, _, rest = cap.partition("  ")
                r1 = cp.add_run(num + "  ")
                set_run_font(r1, BODY_FONT, 8.6, True, ACCENT)
                r2 = cp.add_run(rest if rest else "")
                set_run_font(r2, BODY_FONT, 8.6, False, MUTED)
            emitted_break = False
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            code = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i].rstrip())
                i += 1
            i += 1
            t = doc.add_table(rows=1, cols=1)
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            c = t.cell(0, 0)
            shade(c, "F5F6F7")
            cell_borders(c, "D3D8DD", 4)
            c.text = ""
            for k, cl in enumerate(code):
                par = c.paragraphs[0] if k == 0 else c.add_paragraph()
                par.paragraph_format.space_after = Pt(0)
                par.paragraph_format.line_spacing = 1.05
                r = par.add_run(cl if cl.strip() else " ")
                set_run_font(r, MONO_FONT, 9, color=RGBColor(0x22, 0x2A, 0x33))
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            emitted_break = False
            continue

        # Callouts: "> " note (blue), ">> " term box (green), ">! " pitfall (amber)
        callout = None
        for mark, fill, border in (
            (">> ", "EDF6EE", "8FBF98"),
            (">! ", "FDF3E7", "D9A55F"),
            ("> ", "EEF4F8", "9FBDCC"),
        ):
            if stripped.startswith(mark):
                callout = (mark, fill, border)
                break
        if callout:
            mark, fill, border = callout
            bare = mark.strip()
            block = []
            while i < n and (lines[i].strip().startswith(mark)
                             or lines[i].strip() == bare):
                st = lines[i].strip()
                block.append("" if st == bare else st[len(mark):])
                i += 1
            while block and block[-1] == "":
                block.pop()
            t = doc.add_table(rows=1, cols=1)
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            c = t.cell(0, 0)
            shade(c, fill)
            cell_borders(c, border, 6)
            c.text = ""
            for k, bl in enumerate(block):
                par = c.paragraphs[0] if k == 0 else c.add_paragraph()
                par.paragraph_format.space_after = Pt(2)
                if bl == "":
                    par.paragraph_format.space_after = Pt(0)
                    set_run_font(par.add_run(""), BODY_FONT, 4)
                elif bl.strip().startswith("- "):
                    par.paragraph_format.left_indent = Cm(0.5)
                    par.paragraph_format.first_line_indent = Cm(-0.35)
                    r = par.add_run("• ")
                    set_run_font(r, BODY_FONT, 10)
                    add_inline(par, bl.strip()[2:], 10)
                else:
                    add_inline(par, bl, 10)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            emitted_break = False
            continue

        if stripped.startswith("|"):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                raw = lines[i].strip()
                raw = raw[1:] if raw.startswith("|") else raw
                raw = raw[:-1] if raw.endswith("|") and not raw.endswith("\\|") else raw
                cells = [c.strip().replace("\\|", "|") for c in re.split(r"(?<!\\)\|", raw)]
                rows.append(cells)
                i += 1
            sep_idx = None
            for k, r in enumerate(rows):
                if all(re.fullmatch(r":?-{2,}:?", c) for c in r if c):
                    sep_idx = k
                    break
            header = rows[:sep_idx] if sep_idx is not None else []
            body = rows[sep_idx + 1:] if sep_idx is not None else rows
            ncols = max(len(r) for r in rows)
            t = doc.add_table(rows=0, cols=ncols)
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.autofit = True
            for hr in header:
                cells = t.add_row().cells
                for j in range(ncols):
                    shade(cells[j], "0B4F6C")
                    cell_borders(cells[j], "0B4F6C", 6)
                    par = cells[j].paragraphs[0]
                    par.paragraph_format.space_after = Pt(1)
                    add_inline(par, hr[j] if j < len(hr) else "", 9.5, True,
                               RGBColor(0xFF, 0xFF, 0xFF))
            for ri, br in enumerate(body):
                cells = t.add_row().cells
                for j in range(ncols):
                    if ri % 2 == 1:
                        shade(cells[j], "F4F7F9")
                    cell_borders(cells[j])
                    par = cells[j].paragraphs[0]
                    par.paragraph_format.space_after = Pt(1)
                    add_inline(par, br[j] if j < len(br) else "", 9.5)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        m = re.match(r"^(\s*)([-*])\s+(.*)$", line)
        if m:
            indent = len(m.group(1))
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, m.group(3), 10.5)
            i += 1
            continue

        m = re.match(r"^(\s*)(\d+)[.)]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.85)
            p.paragraph_format.first_line_indent = Cm(-0.85)
            num = p.add_run(f"{m.group(2)}. ")
            set_run_font(num, BODY_FONT, 10.5, True)
            add_inline(p, m.group(3), 10.5)
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        add_inline(p, stripped, 10.5)
        i += 1

    doc.save(out_path)
    print(f"saved: {out_path}")


if __name__ == "__main__":
    build(sys.argv[1], sys.argv[2])
